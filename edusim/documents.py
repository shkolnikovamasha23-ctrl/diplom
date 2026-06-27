from pathlib import Path
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from flask import current_app
from .database import get_db

def _safe_name(prefix, suffix):
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f'{prefix}_{stamp}.{suffix}'

def create_attempt_docx(user, attempt_id):
    db = get_db()
    attempt = db.execute('SELECT a.*, s.title AS scenario_title FROM attempts a JOIN scenarios s ON s.id=a.scenario_id WHERE a.id=?', (attempt_id,)).fetchone()
    events = db.execute('''SELECT st.title AS step_title, c.text AS choice_text, c.feedback, f.original_name,
                                  e.time_spent, e.budget_spent, e.risk_level, e.quality, e.compliance
                           FROM simulation_events e
                           JOIN scenario_steps st ON st.id=e.step_id
                           JOIN scenario_choices c ON c.id=e.choice_id
                           LEFT JOIN uploaded_files f ON f.id=e.uploaded_file_id
                           WHERE e.attempt_id=? ORDER BY e.id''', (attempt_id,)).fetchall()
    answers = db.execute('''SELECT q.text AS question, o.text AS option_text, ans.is_correct
                            FROM answers ans
                            JOIN questions q ON q.id=ans.question_id
                            JOIN options o ON o.id=ans.option_id
                            WHERE ans.attempt_id=?''', (attempt_id,)).fetchall()
    doc = Document()
    doc.add_heading('Отчет о прохождении образовательного симулятора', 0)
    doc.add_paragraph(f"Пользователь: {user['full_name']}")
    doc.add_paragraph(f"Сценарий: {attempt['scenario_title']}")
    doc.add_paragraph(f"Итог: {attempt['ending'] or attempt['status']}")
    doc.add_paragraph(f"Балл: {attempt['score']}")
    doc.add_paragraph(f"Время: {attempt['time_spent']} ч; бюджет: {attempt['budget_spent']} руб.; риск: {attempt['risk_level']}/100; качество: {attempt['quality']}/100; регламент: {attempt['compliance']}/100")
    if attempt['final_report']:
        doc.add_paragraph(attempt['final_report'])
    if events:
        doc.add_heading('История решений', level=1)
        table = doc.add_table(rows=1, cols=5)
        headers = table.rows[0].cells
        headers[0].text = 'Этап'
        headers[1].text = 'Решение'
        headers[2].text = 'Обратная связь'
        headers[3].text = 'Файл'
        headers[4].text = 'Состояние'
        for ev in events:
            row = table.add_row().cells
            row[0].text = ev['step_title']
            row[1].text = ev['choice_text']
            row[2].text = ev['feedback']
            row[3].text = ev['original_name'] or '-'
            row[4].text = f"риск {ev['risk_level']}, качество {ev['quality']}, регламент {ev['compliance']}"
    elif answers:
        doc.add_heading('Ответы тестовой версии', level=1)
        table = doc.add_table(rows=1, cols=3)
        headers = table.rows[0].cells
        headers[0].text = 'Вопрос'
        headers[1].text = 'Выбранный ответ'
        headers[2].text = 'Результат'
        for ans in answers:
            row = table.add_row().cells
            row[0].text = ans['question']
            row[1].text = ans['option_text']
            row[2].text = 'Верно' if ans['is_correct'] else 'Ошибка'
    path = Path(current_app.config['GENERATED_FOLDER']) / _safe_name('attempt_report', 'docx')
    doc.save(path)
    db.execute('INSERT INTO generated_documents(user_id, doc_type, title, file_path, created_at) VALUES (?,?,?,?,?)', (user['id'], 'docx', 'Отчет о прохождении', str(path), datetime.now().isoformat(timespec='seconds')))
    db.commit()
    return path

def create_attempts_xlsx(user):
    db = get_db()
    rows = db.execute('''SELECT a.id, s.title, a.started_at, a.finished_at, a.score, a.status,
                                a.time_spent, a.budget_spent, a.risk_level, a.quality, a.compliance, a.ending
                         FROM attempts a JOIN scenarios s ON s.id=a.scenario_id
                         WHERE a.user_id=? ORDER BY a.id DESC''', (user['id'],)).fetchall()
    wb = Workbook()
    ws = wb.active
    ws.title = 'Результаты'
    ws.append(['ID', 'Сценарий', 'Начало', 'Завершение', 'Балл', 'Статус', 'Время', 'Бюджет', 'Риск', 'Качество', 'Регламент', 'Итог'])
    for row in rows:
        ws.append([row['id'], row['title'], row['started_at'], row['finished_at'], row['score'], row['status'], row['time_spent'], row['budget_spent'], row['risk_level'], row['quality'], row['compliance'], row['ending']])
    path = Path(current_app.config['GENERATED_FOLDER']) / _safe_name('attempts_export', 'xlsx')
    wb.save(path)
    db.execute('INSERT INTO generated_documents(user_id, doc_type, title, file_path, created_at) VALUES (?,?,?,?,?)', (user['id'], 'xlsx', 'Выгрузка результатов', str(path), datetime.now().isoformat(timespec='seconds')))
    db.commit()
    return path
